import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

st.title("📝 Document Formatter")

# Dropdown to choose list style
list_style = st.selectbox("Choose list style:", ["Bullets", "Numbers", "None"])

# Text input
text_input = st.text_area("Paste your rough text here:", height=300)

# Upload a .docx file
uploaded_file = st.file_uploader("Or upload a .docx file", type="docx")

# Process raw text
raw_text = ""

if uploaded_file:
    doc = Document(uploaded_file)
    raw_text = "\n".join([para.text for para in doc.paragraphs])

if text_input:
    raw_text = text_input

# --- Formatter Function ---
def format_text(raw_text, list_style="Bullets"):
    lines = raw_text.strip().split('\n')
    formatted_paragraphs = []
    bullet_mode = False
    number_counter = 1

    for line in lines:
        stripped = line.strip()

        if not stripped:
            bullet_mode = False
            number_counter = 1
            continue

        if stripped.endswith(':'):
            formatted_paragraphs.append(stripped)
            bullet_mode = True
            number_counter = 1
            continue

        formatted_line = stripped[0].upper() + stripped[1:] if stripped else stripped

        if bullet_mode:
            if list_style == "Bullets":
                formatted_paragraphs.append(f"• {formatted_line}")
            elif list_style == "Numbers":
                formatted_paragraphs.append(f"{number_counter}. {formatted_line}")
                number_counter += 1
            else:
                formatted_paragraphs.append(formatted_line)
        else:
            formatted_paragraphs.append(formatted_line)

    return '\n\n'.join(formatted_paragraphs)

# Format and display
if raw_text:
    formatted = format_text(raw_text, list_style=list_style)

    st.subheader("Formatted Preview")
    st.text_area("Formatted Text", value=formatted, height=300)

    st.markdown("*Note: Preview is plain text. Downloaded .docx will be justified and properly formatted.*")

    # Show a nice justified HTML preview
    st.subheader("Justified Preview")
    html_paragraphs = "".join(f"<p>{para}</p>" for para in formatted.split("\n\n"))
    justified_html = f"""
    <div style="text-align: justify; line-height: 1.6; font-size: 16px;">
        {html_paragraphs}
    </div>
    """
    st.markdown(justified_html, unsafe_allow_html=True)

    # Export to docx
    docx_file = io.BytesIO()
    doc = Document()
    for line in formatted.split('\n\n'):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(docx_file)
    docx_file.seek(0)

    st.download_button("Download .docx", docx_file, file_name="formatted.docx")



